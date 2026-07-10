from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "user-guide"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "SpendWise_New_User_Tutorial.docx"

INK = RGBColor(21, 18, 31)
MUTED = RGBColor(112, 106, 128)
VIOLET = RGBColor(105, 79, 255)
CORAL = RGBColor(255, 152, 122)
MINT = RGBColor(107, 211, 166)
LIGHT_BG = "F3EFFB"
LIGHT_FILL = "F8F6FC"
LIGHT_LINE = "DED8EA"
WHITE = "FFFFFF"


SCREENSHOTS = {
    "home": ROOT / "screenshot_insights_dashboard.png",
    "activity": ROOT / "screenshot_insights_tab_active.png",
    "add": ROOT / "screenshots" / "v2-flow-a-save-visible.png",
    "detail": ROOT / "spendwise-dark-detail-sheet.png",
    "insights": ROOT / "screenshot_insights_picker_open.png",
    "menu": ROOT / "spendwise-dark-insights-menu.png",
}


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def prep_images():
    prepared = {}
    for key, src in SCREENSHOTS.items():
        if not src.exists():
            continue
        dest = ASSET_DIR / f"{key}.jpg"
        with Image.open(src) as img:
            img = img.convert("RGB")
            img.thumbnail((780, 1500), Image.Resampling.LANCZOS)
            img.save(dest, quality=86, optimize=True)
        prepared[key] = dest
    return prepared


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LIGHT_LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def paragraph_border_bottom(paragraph, color="D7D0E7", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 26, INK, 0, 8),
        ("Subtitle", 12, MUTED, 0, 16),
        ("Heading 1", 17, INK, 16, 7),
        ("Heading 2", 13.5, VIOLET, 12, 5),
        ("Heading 3", 11.5, INK, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name in ("Title", "Heading 1", "Heading 2", "Heading 3")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.text = "SpendWise New User Tutorial"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Local-first expense tracking guide")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def add_para(doc, text="", style=None, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc, title, body, fill=LIGHT_FILL, accent=VIOLET):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [Inches(6.45)])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = accent
    r.font.size = Pt(10.8)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_ref_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_fill(hdr[idx], LIGHT_BG)
        set_cell_border(hdr[idx])
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = INK
        r.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_fill(cells[idx], WHITE)
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width=2.15, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = align
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def add_two_figures(doc, left, left_caption, right, right_caption):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [Inches(3.2), Inches(3.2)])
    for row in table.rows:
        for cell in row.cells:
            set_cell_fill(cell, WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    for idx, (path, caption) in enumerate([(left, left_caption), (right, right_caption)]):
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.2))
        c = table.cell(1, idx).paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(caption)
        run.italic = True
        run.font.size = Pt(8.3)
        run.font.color.rgb = MUTED
    doc.add_paragraph()


def add_cover(doc, images):
    add_para(doc, "SPENDWISE", bold=True, color=VIOLET, size=11)
    title = add_para(doc, "New User Tutorial & Reference Guide", style="Title")
    title.paragraph_format.space_after = Pt(2)
    add_para(
        doc,
        "A practical guide for recording expenses, reviewing spending, managing accounts and categories, and exporting your yearly transaction data.",
        style="Subtitle",
    )
    meta = add_para(doc, "Generated 26 May 2026 | Currency: MYR (RM) | Data model: local-first on-device storage")
    meta.runs[0].font.color.rgb = MUTED
    meta.runs[0].font.size = Pt(9.5)
    paragraph_border_bottom(meta)
    add_callout(
        doc,
        "Start here",
        "SpendWise is built for fast, private daily tracking: add a transaction, check the month, and move on. There are no accounts to create and no cloud sync prompts.",
        fill=LIGHT_BG,
    )
    if "home" in images and "add" in images:
        add_two_figures(
            doc,
            images["home"],
            "Home shows total balance, net cashflow, heatmap, and recent spending.",
            images["add"],
            "The add sheet captures amount, date, category, account, merchant, and notes.",
        )
    doc.add_page_break()


def build_doc(images):
    doc = Document()
    style_doc(doc)
    add_cover(doc, images)

    add_para(doc, "1. What SpendWise Is For", style="Heading 1")
    add_para(
        doc,
        "SpendWise is a personal Android expense tracker for one person. It stores expenses, categories, accounts, budgets, and appearance settings on your device. The app is tuned for Malaysian Ringgit, so money is shown as RM.",
    )
    add_ref_table(
        doc,
        ["Area", "What it helps you answer"],
        [
            ("Home", "How much money you have, what changed this month, and where spending is concentrated."),
            ("Activity", "Which transactions happened this month, with search and filters for exact lookup."),
            ("Insights", "How a selected year and month are behaving across income, spend, and categories."),
            ("Settings", "Where to manage accounts and categories."),
        ],
        [Inches(1.45), Inches(4.95)],
    )
    add_callout(
        doc,
        "Privacy note",
        "SpendWise does not require sign-in, cloud sync, collaboration, or telemetry. Because the data lives locally, export CSV files periodically if you want an external backup.",
        fill="FFF7F2",
        accent=CORAL,
    )

    add_para(doc, "2. First-Time Setup", style="Heading 1")
    add_para(doc, "Use this short flow the first time you open the app.")
    add_numbers(
        doc,
        [
            "Open SpendWise. The app creates default expense categories and a default Wallet account automatically.",
            "Tap the moon or sun button on Home if you prefer dark mode or light mode.",
            "Tap the Total Balance card, or open Settings > Manage accounts, to add your real wallets, bank accounts, e-wallets, or credit accounts.",
            "Use the plus button to add your first expense or income entry.",
            "Return to Home to confirm your balance, heatmap, recent activity, and category totals update.",
        ],
    )
    add_para(doc, "Default categories", style="Heading 2")
    add_para(
        doc,
        "Fresh installs include Food, Transport, Bills, Shopping, Health, and Salary. Salary is treated as income; the others are treated as expenses. Built-in categories are locked so historical data stays consistent.",
    )

    add_para(doc, "3. Navigation Basics", style="Heading 1")
    add_ref_table(
        doc,
        ["Control", "Use it for"],
        [
            ("Home icon", "Return to the main monthly dashboard."),
            ("Activity icon", "Browse, search, filter, edit, delete, or share transactions."),
            ("Insights icon", "Review yearly cashflow, month trends, category breakdowns, and export CSV."),
            ("Plus button", "Add a new expense or income from anywhere in the main tabs."),
            ("Back button", "Close the current screen or sheet and return to the previous view."),
        ],
        [Inches(1.55), Inches(4.85)],
    )
    add_para(
        doc,
        "The bottom navigation floats over the app. When a sheet or picker is open, the navigation hides so the form has enough space.",
    )

    add_para(doc, "4. Home Tab", style="Heading 1")
    add_para(doc, "Home is the monthly dashboard. It is the fastest place to understand your current money picture.")
    add_bullets(
        doc,
        [
            "Month pill: tap the month near the top left to jump to another month with data.",
            "Search button: opens Activity search.",
            "Theme button: switches light and dark mode.",
            "Settings button: opens account and category management.",
            "Total Balance card: shows active account balance and opens the account list.",
            "Net cashflow: income minus expenses for the selected month.",
            "Spending heatmap: shows which days had heavier spending.",
            "Where you spent: ranks top categories for the month and shows budget progress when a category budget exists.",
            "Recent activity: shows the newest transactions and opens their detail/edit flow.",
        ],
    )
    if "home" in images:
        add_figure(doc, images["home"], "Home dashboard with balance, cashflow, heatmap, and category summary.", width=2.55)

    add_para(doc, "5. Add an Expense or Income", style="Heading 1")
    add_para(doc, "Tap the plus button to open the Add sheet.")
    add_numbers(
        doc,
        [
            "Choose Expense or Income. Existing transactions keep their original type when editing.",
            "Enter the amount. The app accepts digits and one decimal point, with up to two decimal places.",
            "Tap the date chip to choose a transaction date.",
            "Pick a category. Tap New if the category you need does not exist yet.",
            "Pick an account so the app can update the correct balance.",
            "Enter the merchant or source. For expenses this is required; for income it is optional.",
            "Add notes if they will help future-you understand the transaction.",
            "Tap Save expense or Save income.",
        ],
    )
    add_callout(
        doc,
        "Validation rules",
        "The amount must be above RM 0.00, every transaction needs a category and account, expenses need a merchant or description, and dates must be valid.",
        fill=LIGHT_BG,
    )
    if "add" in images:
        add_figure(doc, images["add"], "Adding an expense after selecting category, account, amount, and merchant.", width=2.55)

    add_para(doc, "6. Search, Filter, and Edit Transactions", style="Heading 1")
    add_para(
        doc,
        "Activity is the transaction ledger. It shows entries for the selected month, grouped by day, with totals and filters.",
    )
    add_bullets(
        doc,
        [
            "Search checks merchant, notes, and category names.",
            "All, Expense, and Income filters change the list and the month total.",
            "Account and Category chips open pickers for tighter filtering.",
            "Tap a transaction row to open the detail sheet.",
            "From the detail sheet, you can share a transaction, delete it, or tap Edit transaction.",
        ],
    )
    if "activity" in images and "detail" in images:
        add_two_figures(
            doc,
            images["activity"],
            "Activity lists transactions and supports live search plus filters.",
            images["detail"],
            "Transaction detail shows amount, date, account, note, delete, and edit.",
        )

    add_para(doc, "7. Accounts", style="Heading 1")
    add_para(
        doc,
        "Accounts represent where money lives: cash, bank, e-wallet, or credit. SpendWise calculates current account balance from starting balance plus income minus expenses assigned to that account.",
    )
    add_numbers(
        doc,
        [
            "Open Settings > Manage accounts, or tap the Total Balance card on Home.",
            "Tap Add account.",
            "Enter the account name, type, starting balance, icon, and tile color.",
            "Tap Create account.",
            "To change an account, tap it, edit details, and tap Save changes.",
        ],
    )
    add_callout(
        doc,
        "Archiving accounts",
        "An account can be archived only when no transactions still point to it. If archiving is blocked, move or edit those transactions first. Archived accounts can be restored from the Archived section.",
        fill="FFF7F2",
        accent=CORAL,
    )

    add_para(doc, "8. Categories and Monthly Budgets", style="Heading 1")
    add_para(
        doc,
        "Categories organize transactions and drive the Home and Insights breakdowns. Expense categories can also have an optional monthly budget limit.",
    )
    add_numbers(
        doc,
        [
            "Open Settings > Manage categories, or tap New from the category row while adding a transaction.",
            "Choose whether the category is Expense or Income.",
            "Name it, choose an icon, and choose a tile color.",
            "For expense categories, optionally enter a monthly budget limit.",
            "Tap Save.",
        ],
    )
    add_bullets(
        doc,
        [
            "Built-in categories are locked and cannot be edited or deleted.",
            "Custom categories can be edited.",
            "If you delete a custom category that has transactions, SpendWise asks whether to move the transactions to another category or delete them with the category.",
            "Budget progress appears in category rows; the app nudges around 80 percent and warns at 100 percent.",
        ],
    )

    add_para(doc, "9. Insights and Export", style="Heading 1")
    add_para(
        doc,
        "Insights is the yearly analytics view. It combines earned, spent, net YTD, month-by-month cashflow, and category breakdowns.",
    )
    add_bullets(
        doc,
        [
            "Tap Change beside the year to switch years.",
            "Tap a month bar, or use the arrows in the cashflow card, to inspect a month.",
            "Use Spend and Income to switch the category donut and list.",
            "Open the three-dot menu to export CSV or edit categories.",
            "CSV export includes all transactions for the selected year, regardless of current filters.",
            "In-app PDF export, compare years, and yearly budgets are marked as coming soon.",
        ],
    )
    if "insights" in images and "menu" in images:
        add_two_figures(
            doc,
            images["insights"],
            "Insights shows year, net YTD, cashflow, month selector, and category donut.",
            images["menu"],
            "The Insights menu exports CSV and opens category management.",
        )

    add_para(doc, "10. Month, Year, and Range Pickers", style="Heading 1")
    add_bullets(
        doc,
        [
            "Home and Activity use the month picker to jump between tracked months.",
            "Insights uses a year picker and a month selector inside the cashflow card.",
            "Custom range tools are available for focused review when the app opens a range sheet.",
            "If a screen looks empty, first check the selected month or year, then clear filters.",
        ],
    )

    add_para(doc, "11. Common Problems", style="Heading 1")
    add_ref_table(
        doc,
        ["Problem", "What to try"],
        [
            ("No transactions visible", "Check the selected month, clear search text, and reset filters to All."),
            ("Cannot save expense", "Enter an amount above RM 0.00, choose category/account, and add a merchant or description."),
            ("Cannot archive account", "The account still has transactions. Edit those transactions to another account first."),
            ("Cannot edit built-in category", "Built-in categories are locked. Create a custom category instead."),
            ("CSV export says nothing to export", "Switch Insights to a year that has transactions, then export again."),
            ("Totals look wrong", "Check whether entries were saved as Income or Expense and assigned to the intended account."),
        ],
        [Inches(2.05), Inches(4.35)],
    )

    add_para(doc, "12. Good Habits", style="Heading 1")
    add_bullets(
        doc,
        [
            "Record transactions close to when they happen; merchant suggestions become more useful over time.",
            "Use accounts consistently so balances stay meaningful.",
            "Keep category names simple and stable; reporting is easier when categories are not too granular.",
            "Export a yearly CSV before uninstalling the app, resetting the phone, or moving data elsewhere.",
            "Review Insights monthly to catch spending patterns before they become background noise.",
        ],
    )
    add_callout(
        doc,
        "Quick mental model",
        "Add daily details in Activity, read the month on Home, and use Insights when you want the bigger story.",
        fill=LIGHT_BG,
        accent=VIOLET,
    )
    return doc


def main():
    ensure_dirs()
    images = prep_images()
    doc = build_doc(images)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
